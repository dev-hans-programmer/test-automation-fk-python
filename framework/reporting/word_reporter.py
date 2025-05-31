"""
Word Reporter
Generates Word documents with embedded screenshots
"""

import os
from datetime import datetime
from typing import Dict, Any

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor

from framework.utils.logger import Logger


class WordReporter:
    def __init__(self, config: Dict[str, Any]):
        """Initialize Word Reporter"""
        self.config = config
        self.logger = Logger()
        self.report_dir = config.get('reporting', {}).get('report_directory', './reports')
        
        # Ensure report directory exists
        os.makedirs(self.report_dir, exist_ok=True)
    
    def generate_report(self, execution_data: Dict[str, Any]) -> str:
        """Generate Word report with embedded screenshots"""
        try:
            # Create Word document
            document = Document()
            
            # Add title and header
            self._add_header(document, execution_data)
            
            # Add execution summary
            self._add_execution_summary(document, execution_data)
            
            # Add scenario details
            self._add_scenario_details(document, execution_data)
            
            # Save document
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            execution_id = execution_data.get('execution_id', 'unknown')
            filename = f"test_report_{execution_id}_{timestamp}.docx"
            filepath = os.path.join(self.report_dir, filename)
            
            document.save(filepath)
            
            self.logger.info(f"Word report generated: {filepath}")
            return filepath
            
        except Exception as e:
            self.logger.error(f"Failed to generate Word report: {str(e)}")
            raise
    
    def _add_header(self, document: Document, execution_data: Dict[str, Any]):
        """Add report header"""
        # Title
        title = document.add_heading('Test Automation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Execution info
        execution_id = execution_data.get('execution_id', 'Unknown')
        start_time = execution_data.get('start_time', 'Unknown')
        
        info_para = document.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_para.add_run(f"Execution ID: {execution_id}\n")
        info_para.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
        info_para.add_run(f"Execution Start: {start_time}")
        
        document.add_page_break()
    
    def _add_execution_summary(self, document: Document, execution_data: Dict[str, Any]):
        """Add execution summary section"""
        document.add_heading('Execution Summary', level=1)
        
        # Summary table
        table = document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Metric'
        hdr_cells[1].text = 'Value'
        
        # Data rows
        summary_data = [
            ('Execution ID', execution_data.get('execution_id', 'N/A')),
            ('Status', execution_data.get('status', 'N/A')),
            ('Start Time', execution_data.get('start_time', 'N/A')),
            ('End Time', execution_data.get('end_time', 'N/A')),
            ('Duration (seconds)', str(execution_data.get('duration', 0))),
            ('Total Scenarios', str(execution_data.get('total_scenarios', 0))),
            ('Passed Scenarios', str(execution_data.get('passed_scenarios', 0))),
            ('Failed Scenarios', str(execution_data.get('failed_scenarios', 0))),
            ('Success Rate', f"{self._calculate_success_rate(execution_data)}%")
        ]
        
        for metric, value in summary_data:
            row_cells = table.add_row().cells
            row_cells[0].text = metric
            row_cells[1].text = value
        
        document.add_paragraph()
    
    def _add_scenario_details(self, document: Document, execution_data: Dict[str, Any]):
        """Add detailed scenario information"""
        document.add_heading('Scenario Details', level=1)
        
        for i, scenario in enumerate(execution_data.get('scenarios', [])):
            # Scenario header
            scenario_name = scenario.get('scenario_name', f'Scenario {i+1}')
            status = scenario.get('status', 'unknown')
            
            scenario_heading = document.add_heading(f"{i+1}. {scenario_name}", level=2)
            
            # Status styling
            status_para = document.add_paragraph()
            status_run = status_para.add_run(f"Status: {status.upper()}")
            if status == 'passed':
                status_run.font.color.rgb = RGBColor(0, 128, 0)  # Green
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)  # Red
            status_run.bold = True
            
            # Scenario summary
            scenario_info = [
                f"Duration: {scenario.get('duration', 0)} seconds",
                f"Total Steps: {len(scenario.get('steps', []))}",
                f"Passed Steps: {len([s for s in scenario.get('steps', []) if s.get('status') == 'passed'])}",
                f"Failed Steps: {len([s for s in scenario.get('steps', []) if s.get('status') == 'failed'])}"
            ]
            
            for info in scenario_info:
                document.add_paragraph(info, style='List Bullet')
            
            # Add error if present
            if scenario.get('error'):
                error_para = document.add_paragraph()
                error_para.add_run("Error: ").bold = True
                error_para.add_run(scenario.get('error'))
                error_para.style = 'Intense Quote'
            
            # Add video link if available
            if scenario.get('video_path') and self.config.get('reporting', {}).get('video_embedding', True):
                self._add_video_link(document, scenario.get('video_path'), scenario_name)
            
            # Add step details
            self._add_step_details(document, scenario)
            
            document.add_page_break()
    
    def _add_step_details(self, document: Document, scenario: Dict[str, Any]):
        """Add step details with screenshots"""
        if not scenario.get('steps'):
            return
        
        document.add_heading('Step Details', level=3)
        
        for step in scenario.get('steps', []):
            # Step header
            step_name = step.get('step_name', 'Unknown Step')
            step_status = step.get('status', 'unknown')
            
            step_para = document.add_paragraph()
            step_para.add_run(f"Step {step.get('step_id', '?')}: {step_name}").bold = True
            
            # Step details
            step_info = document.add_paragraph()
            step_info.add_run(f"Action: {step.get('action', 'N/A')}\n")
            step_info.add_run(f"Target: {step.get('target', 'N/A')}\n")
            step_info.add_run(f"Value: {step.get('value', 'N/A')}\n")
            step_info.add_run(f"Duration: {step.get('duration', 0)} seconds\n")
            
            # Status
            status_para = document.add_paragraph()
            status_run = status_para.add_run(f"Status: {step_status.upper()}")
            if step_status == 'passed':
                status_run.font.color.rgb = RGBColor(0, 128, 0)
            else:
                status_run.font.color.rgb = RGBColor(255, 0, 0)
            status_run.bold = True
            
            # Add error if present
            if step.get('error'):
                error_para = document.add_paragraph()
                error_para.add_run("Error: ").bold = True
                error_para.add_run(step.get('error'))
                error_para.style = 'Intense Quote'
            
            # Add screenshot if available
            screenshot_path = step.get('screenshot_path')
            if screenshot_path and os.path.exists(screenshot_path):
                self._add_screenshot(document, screenshot_path, step_name)
            
            document.add_paragraph()  # Add spacing
    
    def _add_screenshot(self, document: Document, screenshot_path: str, caption: str):
        """Add screenshot to document"""
        try:
            # Add screenshot with caption
            screenshot_para = document.add_paragraph()
            screenshot_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image (resize to fit page)
            run = screenshot_para.add_run()
            run.add_picture(screenshot_path, width=Inches(6))
            
            # Add caption
            caption_para = document.add_paragraph()
            caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption_para.add_run(f"Screenshot: {caption}")
            caption_run.italic = True
            caption_run.font.size = 10
            
        except Exception as e:
            self.logger.warning(f"Failed to add screenshot {screenshot_path}: {str(e)}")
            # Add placeholder text instead
            placeholder_para = document.add_paragraph()
            placeholder_para.add_run(f"[Screenshot not available: {caption}]").italic = True
    
    def _add_video_link(self, document: Document, video_path: str, scenario_name: str):
        """Add video link to document"""
        try:
            # Add video section
            video_para = document.add_paragraph()
            video_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add video icon/text
            video_run = video_para.add_run("🎥 Test Execution Video")
            video_run.bold = True
            video_run.font.size = 12
            
            # Add video path/link
            path_para = document.add_paragraph()
            path_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            path_run = path_para.add_run(f"Video Location: {video_path}")
            path_run.italic = True
            path_run.font.size = 10
            
            # Add note about video playback
            note_para = document.add_paragraph()
            note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            note_run = note_para.add_run("(Video can be opened with any media player)")
            note_run.italic = True
            note_run.font.size = 9
            note_run.font.color.rgb = RGBColor(128, 128, 128)
            
        except Exception as e:
            self.logger.warning(f"Failed to add video link {video_path}: {str(e)}")
            # Add placeholder text instead
            placeholder_para = document.add_paragraph()
            placeholder_para.add_run(f"[Video not available: {scenario_name}]").italic = True
    
    def _calculate_success_rate(self, execution_data: Dict[str, Any]) -> float:
        """Calculate success rate percentage"""
        total = execution_data.get('total_scenarios', 0)
        passed = execution_data.get('passed_scenarios', 0)
        
        if total == 0:
            return 0.0
        
        return round((passed / total) * 100, 2)
